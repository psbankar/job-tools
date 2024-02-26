import os
import re
from docx import Document

class ResumeCreator:

    output_folder = "output/"

    latex_config = r"""
    \documentclass[letterpaper,11pt]{article}

    \usepackage{latexsym}
    \usepackage[empty]{fullpage}
    \usepackage{titlesec}
    \usepackage{marvosym}
    \usepackage[usenames,dvipsnames]{color}
    \usepackage{verbatim}
    \usepackage{enumitem}
    \usepackage[hidelinks]{hyperref}
    \usepackage{fancyhdr}
    \usepackage[english]{babel}
    \usepackage{tabularx}
    \input{glyphtounicode}


    %----------FONT OPTIONS----------
    % sans-serif
    % \usepackage[sfdefault]{FiraSans}
    % \usepackage[sfdefault]{roboto}
    % \usepackage[sfdefault]{noto-sans}
    % \usepackage[default]{sourcesanspro}

    % serif
    % \usepackage{CormorantGaramond}
    % \usepackage{charter}


    \pagestyle{fancy}
    \fancyhf{} % clear all header and footer fields
    \fancyfoot{}
    \renewcommand{\headrulewidth}{0pt}
    \renewcommand{\footrulewidth}{0pt}

    % Adjust margins
    \addtolength{\oddsidemargin}{-0.5in}
    \addtolength{\evensidemargin}{-0.5in}
    \addtolength{\textwidth}{1in}
    \addtolength{\topmargin}{-.5in}
    \addtolength{\textheight}{1.0in}

    \urlstyle{same}

    \raggedbottom
    \raggedright
    \setlength{\tabcolsep}{0in}

    % Sections formatting
    \titleformat{\section}{
    \vspace{-4pt}\scshape\raggedright\large
    }{}{0em}{}[\color{black}\titlerule \vspace{-5pt}]

    % Ensure that generate pdf is machine readable/ATS parsable
    \pdfgentounicode=1

    %-------------------------
    % Custom commands
    \newcommand{\resumeItem}[1]{
    \item\small{
        {#1 \vspace{-2pt}}
    }
    }

    \newcommand{\resumeSubheading}[4]{
    \vspace{-2pt}\item
        \begin{tabular*}{0.97\textwidth}[t]{l@{\extracolsep{\fill}}r}
        \textbf{#1} & #2 \\
        \textit{\small#3} & \textit{\small #4} \\
        \end{tabular*}\vspace{-7pt}
    }

    \newcommand{\resumeSubSubheading}[2]{
        \item
        \begin{tabular*}{0.97\textwidth}{l@{\extracolsep{\fill}}r}
        \textit{\small#1} & \textit{\small #2} \\
        \end{tabular*}\vspace{-7pt}
    }

    \newcommand{\resumeProjectHeading}[2]{
        \item
        \begin{tabular*}{0.97\textwidth}{l@{\extracolsep{\fill}}r}
        \small#1 & #2 \\
        \end{tabular*}\vspace{-7pt}
    }

    \newcommand{\resumeSubItem}[1]{\resumeItem{#1}\vspace{-4pt}}

    \renewcommand\labelitemii{$\vcenter{\hbox{\tiny$\bullet$}}$}

    \newcommand{\resumeSubHeadingListStart}{\begin{itemize}[leftmargin=0.15in, label={}]}
    \newcommand{\resumeSubHeadingListEnd}{\end{itemize}}
    \newcommand{\resumeItemListStart}{\begin{itemize}}
    \newcommand{\resumeItemListEnd}{\end{itemize}\vspace{-5pt}}

    %-------------------------------------------
    %%%%%%  RESUME STARTS HERE  %%%%%%%%%%%%%%%%%%%%%%%%%%%%


    \begin{document}
    """

    def generate_latex_header(self, name, phone_no, email, linkedin, github):
        try:
            linkedin = re.sub(r"https://", "", linkedin)
        except:
            pass
        try:
            github = re.sub(r"https://", "", github)
        except:
            pass
        header = (
            r"""
        \begin{center}
            \textbf{\Huge \scshape """
            + name
            + r"""} \\ \vspace{1pt}
            \small """
            + phone_no
            + r""" $|$ \href{mailto:"""
            + email
            + r"""}{\underline{"""
            + email
            + r"""}} $|$ """)
        if linkedin is not None:
            header+= r"""\href{https://"""+ linkedin+ r"""}{\underline{"""+ linkedin+ r"""}} $|$"""
        if github is not None:
            header+= r"""\href{https://"""+ github+ r"""}{\underline{"""+ github + r"""}}"""
        header += r"""
        \end{center}
        """
        return header

    def generate_education(self, entry):

        education = r"""
        \section{Education}
    \resumeSubHeadingListStart"""

        for key, i in entry.items():
            i["date"] = re.sub(r"–", "--", i["date"])
            entry = (
                r"""
                \resumeSubheading
        {"""
                + i["university"]
                + r"""}{"""
                + i["location"]
                + r"""}
        {"""
                + i["degree"]
                + r"""}{"""
                + i["date"]
                + r"""}"""
            )
            education += entry

        education += r"""\resumeSubHeadingListEnd"""
        return education

    def generate_experience(self, entry):
        experience = r"""
        \section{Experience}
    \resumeSubHeadingListStart"""

        for key, i in entry.items():
            i['date'] = re.sub(r"–", "--", i["date"])
            i['description'] = [re.sub(r"%", "\\%", x) for x in i['description']]
            entry = (
                r"""
                \resumeSubheading
        {"""
                + i["company"]
                + r"""}{"""
                + i["location"]
                + r"""}
        {"""
                + i["job_title"]
                + r"""}{"""
                + i["date"]
                + r"""}"""
            )
            experience += entry
            experience += r"""
            \resumeItemListStart
            """
            for point in i["description"]:
                entry =  r"""
                \resumeItem{""" + point + r"""}""" 
                experience += entry
            experience += r"""
            \resumeItemListEnd"""

        experience += r"""\resumeSubHeadingListEnd"""
        return experience

    def generate_projects(self, entry):

        projects = r"""
        \section{Projects}
    \resumeSubHeadingListStart"""

        for key, i in entry.items():
            i['description'] = [re.sub(r"%", "\\%", x) for x in i['description']]
            i["description"] = [re.sub(r"&", "\\&", x) for x in i["description"]]
            i["description"] = [re.sub(r"#", "\\#", x) for x in i["description"]]
            entry = (
                r"""
                \resumeProjectHeading
        {\textbf{"""
                + i["name"]
                + r"""} $|$ \emph{"""
                + ", ".join(i["technologies"])
                + r"""}}{}"""
            )
            projects += entry
            projects += r"""
            \resumeItemListStart
            """
            for point in i["description"]:
                entry =  r"""
                \resumeItem{""" + point + r"""}""" 
                projects += entry
            projects += r"""
            \resumeItemListEnd"""

        projects += r"""\resumeSubHeadingListEnd"""
        return projects

    def generate_skills(self, entry):

        if type(entry) == str:
            entry = entry.split(",")
        skills = (
            r"""
        \section{Skills}
        \begin{itemize}[leftmargin=0.15in, label={}]
        \small{\item{"""
            + ", ".join(entry)
            + r"""}}"""
        )

        skills += r"""\resumeSubHeadingListEnd"""

        skills = re.sub(r"#", "\\#", skills)
        return skills

    def generate_latex_footer(self):
        return r"\end{document}"

    def generate_latex_resume(self, resume_json):
        part1 = self.latex_config
        part2 = self.generate_latex_header(
            resume_json["name"],
            resume_json["phone"],
            resume_json["email"],
            resume_json["url"][2],
            resume_json["url"][1],
        )
        part3 = self.generate_education(resume_json["education"])
        part4 = self.generate_experience(resume_json["experience"])
        part5 = self.generate_projects(resume_json["projects"])
        part6 = self.generate_skills(resume_json["skills"])
        part7 = self.generate_latex_footer()
        return part1 + part2 + part3 + part4 + part5 + part6 + part7

    def generate_pdf(self, resume_json, output_file_name):
        latex_resume = self.generate_latex_resume(resume_json)
        with open("temp.tex", "w") as file:
            file.write(latex_resume)
        os.system(f"pdflatex temp.tex")
        os.replace("temp.pdf", self.output_folder + output_file_name)
        os.remove("temp.tex")
        os.remove("temp.aux")
        os.remove("temp.log")
        os.remove("temp.out")
        return output_file_name

    def generate_docx(self, resume_json, output_file_name):

        doc = Document()
        doc.add_heading(resume_json["name"], 0)
        doc.add_paragraph(resume_json["phone"])
        doc.add_paragraph(resume_json["email"])
        doc.add_paragraph(resume_json["url"][2])
        doc.add_paragraph(resume_json["url"][1])
        doc.add_heading("Education", level=1)
        for key, i in resume_json["education"].items():
            doc.add_paragraph(i["university"] + " - " + i["degree"] + " - " + i["date"])
        doc.add_heading("Experience", level=1)
        for key, i in resume_json["experience"].items():
            doc.add_paragraph(i["company"] + " - " + i["job_title"] + " - " + i["date"])
            for point in i["description"]:
                doc.add_paragraph(point)
        # doc.add_heading("Projects", level=1)
        # for i in resume_json["projects"]:
        #     doc.add_paragraph(i["title"] + " - " + i["date"])
        #     for point in i["description"]:
        #         doc.add_paragraph(point)
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(resume_json["skills"]))
        doc.save(self.output_folder+output_file_name)
        return output_file_name

    def generate_html(self, resume_json, output_file_name):
        html = f"""
        <html>
        <head>
        <title>{resume_json["name"]}</title>
        </head>
        <body>
        <h1>{resume_json["name"]}</h1>
        <p>{resume_json["phone"]}</p>
        <p>{resume_json["email"]}</p>
        <p>{resume_json["url"][2]}</p>
        <p>{resume_json["url"][1]}</p>
        <h2>Education</h2>
        <ul>
        """
        for key,i in resume_json["education"].items():
            html += f"<li>{i['university']} - {i['degree']} - {i['date']}</li>"
        html += f"""
        </ul>
        <h2>Experience</h2>
        <ul>
        """
        for key, i in resume_json["experience"].items():
            html += f"<li>{i['company']} - {i['job_title']} - {i['date']}"
            html += "<ul>"
            for point in i["description"]:
                html += f"<li>{point}</li>"
            html += "</ul></li>"
        # html += f"""
        # </ul>
        # <h2>Projects</h2>
        # <ul>
        # """
        # for i in resume_json["projects"]:
        #     html += f"<li>{i['title']} - {i['date']}"
        #     html += "<ul>"
        #     for point in i["description"]:
        #         html += f"<li>{point}</li>"
        #     html += "</ul></li>"
        html += f"""
        </ul>
        <h2>Skills</h2>
        <ul>
        """
        html += ", ".join([f"<li>{i}</li>" for i in resume_json["skills"]])
        # for i in resume_json["skills"]:
        #     html += f"<li>{i}</li>"
        html += f"""
        </ul>
        </body>
        </html>
        """
        with open(self.output_folder + output_file_name, "w") as file:
            file.write(html)
        return output_file_name

    def generate_markdown(self, resume_json, output_file_name):
        markdown = f"""
        # {resume_json["name"]}
        {resume_json["phone"]}
        {resume_json["email"]}
        {resume_json["url"][2]}
        {resume_json["url"][1]}
        ## Education
        """
        for key, i in resume_json["education"].items():
            markdown += f"* {i['university']} - {i['degree']} - {i['date']}\n"
        markdown += "## Experience\n"
        for key, i in resume_json["experience"].items():
            markdown += f"* {i['company']} - {i['job_title']} - {i['date']}\n"
            for point in i["description"]:
                markdown += f"  * {point}\n"
        # markdown += "## Projects\n"
        # for i in resume_json["projects"]:
        #     markdown += f"* {i['title']} - {i['date']}\n"
        #     for point in i["description"]:
        #         markdown += f"  * {point}\n"
        markdown += "## Skills\n"
        for i in resume_json["skills"]:
            markdown += f"* {i}\n"
        with open(self.output_folder + output_file_name, "w") as file:
            file.write(markdown)
        return output_file_name
